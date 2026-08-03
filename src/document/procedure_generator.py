
"""Generate a professional procedure DOCX from a DocumentBundle."""

from __future__ import annotations

from pathlib import Path

from docx import Document
from docx.document import Document as DocxDocument
from docx.enum.table import WD_CELL_VERTICAL_ALIGNMENT
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.oxml import OxmlElement
from docx.oxml.ns import qn
from docx.shared import Cm, Pt
from docx.table import Table, _Cell
from docx.text.paragraph import Paragraph

from .models import (
    DocumentBundle,
    DocumentOperation,
    DocumentOperationKind,
)


class ProcedureGenerationError(RuntimeError):
    """Raised when the procedure document cannot be generated."""


class ProcedureDocumentGenerator:
    """Populate the company procedure template from a DocumentBundle."""

    def generate(
        self,
        bundle: DocumentBundle,
        template_path: str | Path,
        output_path: str | Path,
    ) -> Path:
        """Generate the procedure document."""

        source_path = Path(template_path)
        destination_path = Path(output_path)

        if not source_path.exists():
            raise FileNotFoundError(
                f"Procedure template not found: {source_path}"
            )

        destination_path.parent.mkdir(
            parents=True,
            exist_ok=True,
        )

        document = Document(source_path)

        self._replace_template_text(
            document=document,
            bundle=bundle,
        )

        self._populate_business_owned_fields(
            document=document,
        )

        self._populate_documents_section(
            document=document,
            bundle=bundle,
        )

        self._populate_unresolved_points(
            document=document,
            bundle=bundle,
        )

        self._compact_presentation_section(
            document=document,
        )

        operations_table = self._find_operations_table(
            document
        )

        self._populate_operations_table(
            table=operations_table,
            operations=bundle.operations,
        )

        internal_controls_table = (
            self._find_internal_controls_table(
                document
            )
        )

        self._populate_internal_controls_table(
            table=internal_controls_table,
            bundle=bundle,
        )

        business_rules_table = (
            self._find_business_rules_table(
                document
            )
        )

        print(
            "Business rules found:",
            len(bundle.procedure.business_rules),
        )

        self._populate_business_rules_table(
            table=business_rules_table,
            bundle=bundle,
        )

        document.save(destination_path)

        return destination_path

    def _replace_template_text(
        self,
        document: DocxDocument,
        bundle: DocumentBundle,
    ) -> None:
        """Replace known text placeholders throughout the document."""

        replacements = {
            "Traitement des demandes d’achat": bundle.metadata.title,
            "Cette procédure décrit la démarche à suivre pour la gestion "
            "des demandes d’achat.": (
                bundle.procedure.purpose
                or (
                    "Cette procédure décrit les opérations du processus "
                    f"« {bundle.metadata.title} »."
                )
            ),
            "Date de dernière modif. :": (
                "Date de dernière modif. : "
                + (bundle.metadata.last_modified_date or "À compléter")
            ),
            "Date de dernière modif.\xa0:": (
                "Date de dernière modif.\xa0: "
                + (bundle.metadata.last_modified_date or "À compléter")
            ),
        }

        if bundle.metadata.domain:
            replacements["Gestion des achats"] = (
                bundle.metadata.domain
            )

        if bundle.metadata.reference:
            replacements["CNI-ACH-003"] = (
                bundle.metadata.reference
            )

        if bundle.metadata.creation_date:
            replacements["01/09/2022"] = (
                bundle.metadata.creation_date
            )

        seen_paragraphs: set[object] = set()

        for paragraph in self._iter_all_paragraphs(
            document
        ):
            element_key = paragraph._p

            if element_key in seen_paragraphs:
                continue

            seen_paragraphs.add(element_key)
            self._replace_text_in_paragraph(
                paragraph=paragraph,
                replacements=replacements,
            )

    @classmethod
    def _populate_business_owned_fields(
        cls,
        *,
        document: DocxDocument,
    ) -> None:
        """Mark administrative fields that require business ownership."""

        expected_headers = {
            ("nom", "signature"),
            ("nom", "fonction", "date", "signature"),
            ("nom", "fonction"),
        }

        for table in document.tables:
            if len(table.rows) < 2:
                continue

            headers = tuple(
                cls._normalize_text(cell.text)
                for cell in table.rows[0].cells
            )

            if headers not in expected_headers:
                continue

            first_data_row = table.rows[1]

            if any(cell.text.strip() for cell in first_data_row.cells):
                continue

            cls._set_cell_text(
                cell=first_data_row.cells[0],
                text="À compléter par le métier",
                alignment=WD_ALIGN_PARAGRAPH.LEFT,
            )

    @classmethod
    def _populate_documents_section(
        cls,
        *,
        document: DocxDocument,
        bundle: DocumentBundle,
    ) -> None:
        """Populate Section C even when the template has no table."""

        heading = cls._find_paragraph_containing(
            document,
            "Documents / Etats utilisés",
        )

        if heading is None:
            return

        if bundle.documents:
            items = []

            for document_item in bundle.documents:
                operation_numbers = [
                    str(operation.number)
                    for operation in bundle.operations
                    if operation.bpmn_element_id
                    in {
                        *document_item.produced_by_operation_ids,
                        *document_item.consumed_by_operation_ids,
                    }
                ]
                reference = (
                    f" (opérations {', '.join(operation_numbers)})"
                    if operation_numbers
                    else ""
                )
                items.append(f"{document_item.name}{reference}")

            text = " ; ".join(items) + "."
        else:
            text = (
                "Non renseigné dans le modèle BPMN — "
                "à compléter par le métier."
            )

        paragraph = cls._insert_paragraph_after(heading)
        cls._configure_paragraph(
            paragraph,
            alignment=WD_ALIGN_PARAGRAPH.LEFT,
            space_after=Pt(8),
        )
        cls._add_run(paragraph=paragraph, text=text, italic=True)

    @classmethod
    def _populate_unresolved_points(
        cls,
        *,
        document: DocxDocument,
        bundle: DocumentBundle,
    ) -> None:
        """Expose unresolved business parameters instead of hiding them."""

        points = bundle.specification.unresolved_points

        if not points:
            return

        purpose = cls._find_paragraph_containing(
            document,
            bundle.procedure.purpose or "",
        )

        if purpose is None:
            return

        paragraph = cls._insert_paragraph_after(purpose)
        cls._configure_paragraph(
            paragraph,
            alignment=WD_ALIGN_PARAGRAPH.LEFT,
            space_before=Pt(4),
            space_after=Pt(8),
        )
        cls._add_run(
            paragraph=paragraph,
            text="Points à confirmer : ",
            bold=True,
        )
        cls._add_run(
            paragraph=paragraph,
            text=" ; ".join(points),
            italic=True,
        )


    @classmethod
    def _compact_presentation_section(
        cls,
        *,
        document: DocxDocument,
    ) -> None:
        """Remove template spacer paragraphs before the next section break."""

        paragraphs = list(document.paragraphs)
        anchor_index = next(
            (
                index
                for index, paragraph in enumerate(paragraphs)
                if "se référer à" in paragraph.text.casefold()
            ),
            None,
        )

        if anchor_index is None:
            return

        removable: list[Paragraph] = []

        for paragraph in paragraphs[anchor_index + 1 :]:
            paragraph_properties = paragraph._p.pPr
            has_section_break = (
                paragraph_properties is not None
                and paragraph_properties.sectPr is not None
            )

            if has_section_break:
                break

            if paragraph.text.strip():
                removable.clear()
                continue

            removable.append(paragraph)

        for paragraph in removable:
            parent = paragraph._p.getparent()
            if parent is not None:
                parent.remove(paragraph._p)

    @classmethod
    def _find_paragraph_containing(
        cls,
        document: DocxDocument,
        text: str,
    ) -> Paragraph | None:
        normalized_target = cls._normalize_text(text)

        if not normalized_target:
            return None

        for paragraph in cls._iter_all_paragraphs(document):
            if normalized_target in cls._normalize_text(paragraph.text):
                return paragraph

        return None

    @staticmethod
    def _insert_paragraph_after(
        paragraph: Paragraph,
    ) -> Paragraph:
        new_element = OxmlElement("w:p")
        paragraph._p.addnext(new_element)
        return Paragraph(new_element, paragraph._parent)

    @staticmethod
    def _format_actor_name(
        actor_name: str | None,
    ) -> str:
        return " ".join(str(actor_name or "").split()).strip()

    @classmethod
    def _find_operations_table(
        cls,
        document: DocxDocument,
    ) -> Table:
        """Find the operations table from its column headers."""

        for table in document.tables:
            if not table.rows:
                continue

            headers = [
                cls._normalize_text(cell.text)
                for cell in table.rows[0].cells
            ]

            if len(headers) < 4:
                continue

            is_operations_table = (
                headers[0] == "intervenant"
                and headers[1] in {
                    "n° op",
                    "nº op",
                    "no op",
                }
                and "description de l'opération" in headers[2]
                and headers[3] == "document"
            )

            if is_operations_table:
                return table

        detected_headers = [
            [
                cls._normalize_text(cell.text)
                for cell in table.rows[0].cells
            ]
            for table in document.tables
            if table.rows
        ]

        raise ProcedureGenerationError(
            "Could not find the operations table. "
            f"Detected table headers: {detected_headers}"
        )

    @classmethod
    def _find_business_rules_table(
        cls,
        document: DocxDocument,
    ) -> Table:
        """Find the business-rules table from its column headers."""

        for table in document.tables:
            if not table.rows:
                continue

            headers = [
                cls._normalize_text(cell.text)
                for cell in table.rows[0].cells
            ]

            if len(headers) < 2:
                continue

            first_header = headers[0]
            second_header = headers[1]

            is_business_rules_table = (
                first_header in {
                    "réf. op",
                    "ref. op",
                    "réf op",
                    "ref op",
                }
                and "règle de gestion" in second_header
            )

            if is_business_rules_table:
                return table

        detected_headers = [
            [
                cls._normalize_text(cell.text)
                for cell in table.rows[0].cells
            ]
            for table in document.tables
            if table.rows
        ]

        raise ProcedureGenerationError(
            "Could not find the business-rules table. "
            f"Detected table headers: {detected_headers}"
        )

    @classmethod
    def _find_internal_controls_table(
        cls,
        document: DocxDocument,
    ) -> Table:
        """Find the internal-controls table from its headers."""

        for table in document.tables:
            if not table.rows:
                continue

            headers = [
                cls._normalize_text(cell.text)
                for cell in table.rows[0].cells
            ]

            if len(headers) < 6:
                continue

            is_internal_controls_table = (
                headers[0] in {
                    "réf. op",
                    "ref. op",
                    "réf op",
                    "ref op",
                }
                and "point de contrôle" in headers[1]
                and "description du contrôle" in headers[2]
                and "objectif du contrôle" in headers[3]
                and "support du contrôle" in headers[4]
                and headers[5] == "responsable"
            )

            if is_internal_controls_table:
                return table

        detected_headers = [
            [
                cls._normalize_text(cell.text)
                for cell in table.rows[0].cells
            ]
            for table in document.tables
            if table.rows
        ]

        raise ProcedureGenerationError(
            "Could not find the internal-controls table. "
            f"Detected table headers: {detected_headers}"
        )

    @staticmethod
    def _build_branch_conditions(
        operations: list[DocumentOperation],
    ) -> dict[str, tuple[str, str]]:
        """Build labels for true branch-entry operations.

        A branch label is rendered on the operation that starts that branch.
        It is not rendered on an operation identified as the common
        continuation of the same decision gateway.
        """

        operations_by_id = {
            operation.bpmn_element_id: operation
            for operation in operations
        }

        conditions_by_target: dict[
            str,
            list[tuple[str, str]],
        ] = {}

        for source_operation in operations:
            for branch in source_operation.branches:
                target = operations_by_id.get(
                    branch.target_operation_id
                )

                if target is None:
                    continue

                # Loop-back branches are represented with the decision,
                # not as the normal entry condition of the target operation.
                if getattr(
                    branch,
                    "is_loop_back",
                    False,
                ):
                    continue

                # Do not attach one branch label to an operation where
                # the branches of that same gateway converge.
                if (
                    target.is_common_continuation
                    and branch.gateway_id
                    in target.convergence_gateway_ids
                ):
                    continue

                branch_label = (
                    branch.label
                    or branch.condition
                )

                if not branch_label:
                    continue

                condition = (
                    branch.gateway_name or "Décision",
                    branch_label,
                )

                conditions_by_target.setdefault(
                    branch.target_operation_id,
                    [],
                ).append(condition)

        result: dict[str, tuple[str, str]] = {}

        for target_id, conditions in conditions_by_target.items():
            unique_conditions = list(
                dict.fromkeys(conditions)
            )

            # Several different conditions attached to one target would
            # be misleading, so only render an unambiguous entry condition.
            if len(unique_conditions) == 1:
                result[target_id] = unique_conditions[0]

        return result

    @classmethod
    def _populate_operations_table(
        cls,
        table: Table,
        operations: list[DocumentOperation],
    ) -> None:
        """
        Clear template operation rows and insert one professional row
        per operation.
        """

        while len(table.rows) > 1:
            row = table.rows[-1]
            table._tbl.remove(row._tr)

        cls._set_operations_table_widths(
            table
        )

        cls._repeat_table_header(
            table.rows[0]
        )

        branch_condition_by_target_id = (
            cls._build_branch_conditions(
                operations
            )
        )

        for operation in operations:
            row = table.add_row()

            cls._prevent_row_split(
                row
            )

            actor_text = cls._format_actor_name(
                operation.actor_name
            )

            document_names = list(
                dict.fromkeys(
                    [
                        *operation.input_document_names,
                        *operation.output_document_names,
                    ]
                )
            )

            cls._set_cell_text(
                cell=row.cells[0],
                text=actor_text,
                alignment=WD_ALIGN_PARAGRAPH.LEFT,
            )

            cls._set_cell_text(
                cell=row.cells[1],
                text=str(operation.number),
                alignment=WD_ALIGN_PARAGRAPH.CENTER,
            )

            if (
                operation.element_kind
                == DocumentOperationKind.SUBPROCESS
            ):
                merged = row.cells[2].merge(
                    row.cells[3]
                )

                cls._render_operation_cell(
                    cell=merged,
                    operation=operation,
                    branch_condition=(
                        branch_condition_by_target_id.get(
                            operation.bpmn_element_id
                        )
                    ),
                )
            else:
                cls._render_operation_cell(
                    cell=row.cells[2],
                    operation=operation,
                    branch_condition=(
                        branch_condition_by_target_id.get(
                            operation.bpmn_element_id
                        )
                    ),
                )

                cls._render_documents_cell(
                    cell=row.cells[3],
                    document_names=document_names,
                )

            row.cells[0].vertical_alignment = (
                WD_CELL_VERTICAL_ALIGNMENT.TOP
            )
            row.cells[1].vertical_alignment = (
                WD_CELL_VERTICAL_ALIGNMENT.CENTER
            )
            row.cells[2].vertical_alignment = (
                WD_CELL_VERTICAL_ALIGNMENT.TOP
            )
            row.cells[3].vertical_alignment = (
                WD_CELL_VERTICAL_ALIGNMENT.TOP
            )

    @classmethod
    def _render_operation_cell(
        cls,
        cell: _Cell,
        operation: DocumentOperation,
        branch_condition: tuple[str, str] | None = None,
    ) -> None:
        """Render description, notes and branches with separate styles."""

        cls._clear_cell(cell)

        description_paragraph = cell.paragraphs[0]

        if operation.is_common_continuation:
            cls._configure_paragraph(
                description_paragraph,
                alignment=WD_ALIGN_PARAGRAPH.LEFT,
                space_after=Pt(1),
            )
            cls._add_run(
                paragraph=description_paragraph,
                text="Convergence",
                bold=True,
            )

            convergence_paragraph = cell.add_paragraph()
            cls._configure_paragraph(
                convergence_paragraph,
                alignment=WD_ALIGN_PARAGRAPH.LEFT,
                space_after=Pt(3),
            )
            cls._add_run(
                paragraph=convergence_paragraph,
                text=(
                    "Les différents scénarios se rejoignent "
                    "avant cette opération."
                ),
                italic=True,
            )

            for branch in operation.direct_convergence_branches:
                branch_label = (
                    branch.label
                    or branch.condition
                    or "Branche non libellée"
                )
                branch_paragraph = cell.add_paragraph()
                cls._configure_paragraph(
                    branch_paragraph,
                    alignment=WD_ALIGN_PARAGRAPH.LEFT,
                    space_after=Pt(3),
                )
                cls._add_run(
                    paragraph=branch_paragraph,
                    text=(
                        f"La branche « {branch_label} » "
                        "accède directement à cette opération."
                    ),
                    italic=True,
                )

            description_paragraph = cell.add_paragraph()

        if branch_condition is not None:
            gateway_name, branch_label = branch_condition

            cls._configure_paragraph(
                description_paragraph,
                alignment=WD_ALIGN_PARAGRAPH.LEFT,
                space_after=Pt(3),
            )

            cls._add_run(
                paragraph=description_paragraph,
                text=f"{gateway_name} — {branch_label}",
            )

            description_paragraph = cell.add_paragraph()

        cls._configure_paragraph(
            description_paragraph,
            alignment=WD_ALIGN_PARAGRAPH.LEFT,
            space_after=Pt(4),
        )

        cls._add_run(
            paragraph=description_paragraph,
            text=operation.description.strip(),
        )

        non_incorporated_notes = [
            note.text.strip()
            for note in operation.notes
            if (
                note.text.strip()
                and not note.incorporated_in_description
            )
        ]

        for note_text in non_incorporated_notes:
            note_paragraph = cell.add_paragraph()

            cls._configure_paragraph(
                note_paragraph,
                alignment=WD_ALIGN_PARAGRAPH.LEFT,
                space_before=Pt(3),
                space_after=Pt(2),
            )

            cls._add_run(
                paragraph=note_paragraph,
                text="Remarque : ",
                bold=True,
                italic=True,
            )

            cls._add_run(
                paragraph=note_paragraph,
                text=note_text,
                italic=True,
            )

        if operation.branches:
            cls._render_decision_question_only(
                cell=cell,
                operation=operation,
            )

    @classmethod
    def _render_decision_question_only(
        cls,
        cell: _Cell,
        operation: DocumentOperation,
    ) -> None:
        """Render only the BPMN gateway question for a decision."""

        gateway_name = next(
            (
                branch.gateway_name
                for branch in operation.branches
                if branch.gateway_name
            ),
            "Décision",
        )

        title_paragraph = cell.add_paragraph()

        cls._configure_paragraph(
            title_paragraph,
            alignment=WD_ALIGN_PARAGRAPH.LEFT,
            space_before=Pt(6),
            space_after=Pt(1),
        )

        cls._add_run(
            paragraph=title_paragraph,
            text="Décision",
            bold=True,
        )

        question_paragraph = cell.add_paragraph()

        cls._configure_paragraph(
            question_paragraph,
            alignment=WD_ALIGN_PARAGRAPH.LEFT,
            space_after=Pt(4),
        )

        cls._add_run(
            paragraph=question_paragraph,
            text=gateway_name,
            bold=True,
            italic=True,
        )

        for branch in operation.branches:
            if not branch.is_loop_back:
                continue

            loop_paragraph = cell.add_paragraph()
            cls._configure_paragraph(
                loop_paragraph,
                alignment=WD_ALIGN_PARAGRAPH.LEFT,
                space_after=Pt(2),
            )
            label = branch.label or branch.condition or "Retour"
            target = (
                branch.target_operation_name
                or "l’opération précédente"
            )
            cls._add_run(
                paragraph=loop_paragraph,
                text=f"{label} : ",
                bold=True,
            )
            cls._add_run(
                paragraph=loop_paragraph,
                text=f"retour à l’activité « {target} ».",
            )

    @staticmethod
    def _find_table_after_heading(
        document: DocxDocument,
        heading_text: str,
    ) -> Table | None:
        """
        Find the first table appearing after a matching paragraph.

        This uses the XML order because document.tables alone does not expose
        surrounding paragraph relationships.
        """

        body_elements = list(
            document.element.body.iterchildren()
        )

        heading_found = False

        for element in body_elements:
            tag_name = element.tag.rsplit("}", 1)[-1]

            if tag_name == "p":
                paragraph = Paragraph(
                    element,
                    document,
                )

                if (
                    heading_text.lower()
                    in paragraph.text.strip().lower()
                ):
                    heading_found = True

            elif tag_name == "tbl" and heading_found:
                return Table(
                    element,
                    document,
                )

        return None

    @staticmethod
    def _populate_documents_table(
        table: Table,
        bundle: DocumentBundle,
    ) -> None:
        """Populate the documents/states table when one exists."""

        if not table.rows:
            return

        while len(table.rows) > 1:
            row = table.rows[-1]
            table._tbl.remove(row._tr)

        if not bundle.documents:
            row = table.add_row()

            for cell in row.cells:
                ProcedureDocumentGenerator._set_cell_text(
                    cell=cell,
                    text="N/A",
                )

            return

        for document_item in bundle.documents:
            row = table.add_row()

            operation_numbers = [
                str(operation.number)
                for operation in bundle.operations
                if (
                    operation.bpmn_element_id
                    in {
                        *document_item.produced_by_operation_ids,
                        *document_item.consumed_by_operation_ids,
                    }
                )
            ]

            values = [
                ", ".join(operation_numbers),
                document_item.name,
            ]

            for index, cell in enumerate(row.cells):
                value = (
                    values[index]
                    if index < len(values)
                    else ""
                )

                ProcedureDocumentGenerator._set_cell_text(
                    cell=cell,
                    text=value,
                )

    @classmethod
    def _populate_internal_controls_table(
        cls,
        table: Table,
        bundle: DocumentBundle,
    ) -> None:
        """Populate Section D or clear template examples."""

        if not table.rows:
            raise ProcedureGenerationError(
                "The internal-controls table has no header row."
            )

        while len(table.rows) > 1:
            row = table.rows[-1]
            table._tbl.remove(row._tr)

        controls = bundle.procedure.internal_controls

        if not controls:
            empty_row = table.add_row()
            merged = empty_row.cells[0]

            for cell in empty_row.cells[1:]:
                merged = merged.merge(cell)

            cls._set_cell_text(
                cell=merged,
                text=(
                    "Non renseigné dans le modèle BPMN — "
                    "à compléter par le métier."
                ),
                alignment=WD_ALIGN_PARAGRAPH.LEFT,
            )

            return

        for control in controls:
            row = table.add_row()

            cls._prevent_row_split(row)

            values = [
                str(
                    control.get(
                        "operation_number",
                        "",
                    )
                ),
                control.get(
                    "control_point",
                    "",
                ),
                control.get(
                    "description",
                    "",
                ),
                control.get(
                    "objective",
                    "",
                ),
                control.get(
                    "support",
                    "",
                ),
                control.get(
                    "responsible",
                    "",
                ),
            ]

            for index, cell in enumerate(
                row.cells
            ):
                cls._set_cell_text(
                    cell=cell,
                    text=(
                        values[index]
                        if index < len(values)
                        else ""
                    ),
                    alignment=(
                        WD_ALIGN_PARAGRAPH.CENTER
                        if index == 0
                        else WD_ALIGN_PARAGRAPH.LEFT
                    ),
                )

                cell.vertical_alignment = (
                    WD_CELL_VERTICAL_ALIGNMENT.TOP
                )

    @classmethod
    def _populate_business_rules_table(
        cls,
        table: Table,
        bundle: DocumentBundle,
    ) -> None:
        """Populate Section E from BPMN annotations."""

        if not table.rows:
            raise ProcedureGenerationError(
                "The business-rules table has no header row."
            )

        # Remove all template/example rows.
        while len(table.rows) > 1:
            row = table.rows[-1]
            table._tbl.remove(row._tr)

        rules = bundle.procedure.business_rules

        if not rules:
            row = table.add_row()

            cls._set_cell_text(
                cell=row.cells[0],
                text="",
                alignment=WD_ALIGN_PARAGRAPH.CENTER,
            )

            cls._set_cell_text(
                cell=row.cells[1],
                text="Aucune règle de gestion identifiée.",
                alignment=WD_ALIGN_PARAGRAPH.LEFT,
            )

            return

        for rule in rules:
            row = table.add_row()

            cls._prevent_row_split(row)

            cls._set_cell_text(
                cell=row.cells[0],
                text=str(rule.operation_number),
                alignment=WD_ALIGN_PARAGRAPH.CENTER,
            )

            cls._set_cell_text(
                cell=row.cells[1],
                text=rule.text.strip(),
                alignment=WD_ALIGN_PARAGRAPH.LEFT,
            )

            row.cells[0].vertical_alignment = (
                WD_CELL_VERTICAL_ALIGNMENT.CENTER
            )

            row.cells[1].vertical_alignment = (
                WD_CELL_VERTICAL_ALIGNMENT.TOP
            )

    @classmethod
    def _render_documents_cell(
        cls,
        cell: _Cell,
        document_names: list[str],
    ) -> None:
        """Render associated documents, one per paragraph."""

        cls._clear_cell(
            cell
        )

        if not document_names:
            return

        for index, document_name in enumerate(
            document_names
        ):
            paragraph = (
                cell.paragraphs[0]
                if index == 0
                else cell.add_paragraph()
            )

            cls._configure_paragraph(
                paragraph,
                alignment=WD_ALIGN_PARAGRAPH.LEFT,
                space_after=Pt(2),
            )

            cls._add_run(
                paragraph=paragraph,
                text=document_name,
            )

    @classmethod
    def _set_cell_text(
        cls,
        cell: _Cell,
        text: str,
        alignment: WD_ALIGN_PARAGRAPH = (
            WD_ALIGN_PARAGRAPH.LEFT
        ),
    ) -> None:
        """Set plain cell text with consistent formatting."""

        cls._clear_cell(
            cell
        )

        lines = text.splitlines() or [""]

        for index, line in enumerate(lines):
            paragraph = (
                cell.paragraphs[0]
                if index == 0
                else cell.add_paragraph()
            )

            cls._configure_paragraph(
                paragraph,
                alignment=alignment,
                space_after=Pt(0),
            )

            cls._add_run(
                paragraph=paragraph,
                text=line,
            )

    @staticmethod
    def _clear_cell(
        cell: _Cell,
    ) -> None:
        """Remove text while preserving one usable paragraph."""

        cell.text = ""

        paragraph = cell.paragraphs[0]

        for run in paragraph.runs:
            run.text = ""

    @staticmethod
    def _configure_paragraph(
        paragraph: Paragraph,
        *,
        alignment: WD_ALIGN_PARAGRAPH,
        space_before: Pt = Pt(0),
        space_after: Pt = Pt(0),
        left_indent: Cm | None = None,
    ) -> None:
        paragraph.alignment = alignment

        paragraph.paragraph_format.space_before = (
            space_before
        )
        paragraph.paragraph_format.space_after = (
            space_after
        )
        paragraph.paragraph_format.line_spacing = 1.05

        if left_indent is not None:
            paragraph.paragraph_format.left_indent = (
                left_indent
            )

    @staticmethod
    def _add_run(
        paragraph: Paragraph,
        text: str,
        *,
        bold: bool = False,
        italic: bool = False,
    ) -> None:
        run = paragraph.add_run(
            text
        )

        run.font.name = "Arial"
        run.font.size = Pt(8.5)
        run.bold = bold
        run.italic = italic

    @staticmethod
    def _lowercase_first_letter(
        value: str,
    ) -> str:
        value = value.strip()

        if not value:
            return value

        return value[0].lower() + value[1:]

    @staticmethod
    def _set_operations_table_widths(
        table: Table,
    ) -> None:
        """Set balanced widths for the four operation columns."""

        widths = [
            Cm(5.0),
            Cm(1.2),
            Cm(9.8),
            Cm(2.4),
        ]

        table.autofit = False

        grid = table._tbl.tblGrid
        grid_columns = list(grid.gridCol_lst)

        for grid_column, width in zip(
            grid_columns,
            widths,
            strict=False,
        ):
            grid_column.set(
                qn("w:w"),
                str(int(width.twips)),
            )

        for row in table.rows:
            for cell, width in zip(
                row.cells,
                widths,
                strict=False,
            ):
                cell.width = width
                cell_width = cell._tc.get_or_add_tcPr().get_or_add_tcW()
                cell_width.set(qn("w:w"), str(int(width.twips)))
                cell_width.set(qn("w:type"), "dxa")

    @staticmethod
    def _repeat_table_header(
        row,
    ) -> None:
        table_row_properties = (
            row._tr.get_or_add_trPr()
        )

        repeat_header = OxmlElement(
            "w:tblHeader"
        )

        repeat_header.set(
            qn("w:val"),
            "true",
        )

        table_row_properties.append(
            repeat_header
        )

    @staticmethod
    def _prevent_row_split(
        row,
    ) -> None:
        row_properties = (
            row._tr.get_or_add_trPr()
        )

        cannot_split = OxmlElement(
            "w:cantSplit"
        )

        row_properties.append(
            cannot_split
        )

    @staticmethod
    def _replace_text_in_paragraph(
        paragraph: Paragraph,
        replacements: dict[str, str],
    ) -> None:
        """
        Replace text while preserving the formatting of the first run.

        This is suitable for simple template labels and titles.
        """

        original_text = paragraph.text

        updated_text = original_text

        for source, destination in replacements.items():
            if destination in updated_text:
                continue

            updated_text = updated_text.replace(
                source,
                destination,
            )

        if updated_text == original_text:
            return

        if not paragraph.runs:
            paragraph.add_run(updated_text)
            return

        first_run = paragraph.runs[0]

        for run in paragraph.runs:
            run.text = ""

        first_run.text = updated_text

    @classmethod
    def _iter_all_paragraphs(
        cls,
        document: DocxDocument,
    ):
        """Yield paragraphs from body, tables, headers and footers."""

        for paragraph in document.paragraphs:
            yield paragraph

        for table in document.tables:
            yield from cls._iter_table_paragraphs(
                table
            )

        for section in document.sections:
            for paragraph in section.header.paragraphs:
                yield paragraph

            for table in section.header.tables:
                yield from cls._iter_table_paragraphs(
                    table
                )

            for paragraph in section.footer.paragraphs:
                yield paragraph

            for table in section.footer.tables:
                yield from cls._iter_table_paragraphs(
                    table
                )

    @classmethod
    def _iter_table_paragraphs(
        cls,
        table: Table,
    ):
        for row in table.rows:
            for cell in row.cells:
                for paragraph in cell.paragraphs:
                    yield paragraph

                for nested_table in cell.tables:
                    yield from cls._iter_table_paragraphs(
                        nested_table
                    )

    @staticmethod
    def _normalize_text(
        value: str,
    ) -> str:
        normalized = (
            value
            .replace("\xa0", " ")
            .replace("\u2019", "'")
            .replace("\u2018", "'")
            .replace("\u201B", "'")
            .replace("\u2032", "'")
            .strip()
            .lower()
        )

        return " ".join(normalized.split())

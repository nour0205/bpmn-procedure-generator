
"""Generate a professional specification DOCX."""

from __future__ import annotations

from pathlib import Path

from docx import Document
from docx.document import Document as DocxDocument
from docx.enum.section import WD_SECTION
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.oxml import OxmlElement
from docx.oxml.ns import qn
from docx.shared import Cm, Pt, RGBColor
from docx.text.paragraph import Paragraph

from .models import DocumentBundle
from .narrative_models import GeneratedNarrative


class SpecificationGenerationError(RuntimeError):
    """Raised when the specification cannot be generated."""


class SpecificationDocumentGenerator:
    """Generate the specification narrative document."""

    def generate(
        self,
        *,
        bundle: DocumentBundle,
        narrative: GeneratedNarrative,
        output_path: str | Path,
    ) -> Path:
        """Generate a DOCX containing the process narrative."""

        self._validate_inputs(
            bundle=bundle,
            narrative=narrative,
        )

        destination_path = Path(
            output_path
        )

        destination_path.parent.mkdir(
            parents=True,
            exist_ok=True,
        )

        document = Document()

        self._configure_document(
            document
        )

        self._add_header(
            document=document,
            bundle=bundle,
        )

        self._add_title(
            document=document,
            title=bundle.metadata.title,
        )

        self._add_metadata_block(
            document=document,
            bundle=bundle,
        )

        self._add_narrative_section(
            document=document,
            paragraphs=narrative.paragraphs,
        )

        document.save(
            destination_path
        )

        return destination_path

    @staticmethod
    def _validate_inputs(
        *,
        bundle: DocumentBundle,
        narrative: GeneratedNarrative,
    ) -> None:
        """Ensure that both files describe the same process."""

        if (
            bundle.metadata.process_id
            != narrative.process_id
        ):
            raise SpecificationGenerationError(
                "The DocumentBundle and generated narrative "
                "use different process IDs."
            )

        if not narrative.paragraphs:
            raise SpecificationGenerationError(
                "The generated narrative contains no paragraphs."
            )

    @classmethod
    def _configure_document(
        cls,
        document: DocxDocument,
    ) -> None:
        """Configure page layout and shared Word styles."""

        section = document.sections[0]

        section.top_margin = Cm(2.2)
        section.bottom_margin = Cm(2.0)
        section.left_margin = Cm(2.4)
        section.right_margin = Cm(2.4)

        normal_style = document.styles["Normal"]

        normal_style.font.name = "Arial"
        normal_style.font.size = Pt(11)

        normal_style.element.rPr.rFonts.set(
            qn("w:eastAsia"),
            "Arial",
        )

        title_style = document.styles["Title"]

        title_style.font.name = "Arial"
        title_style.font.size = Pt(16)
        title_style.font.bold = True
        title_style.font.color.rgb = RGBColor(
            31,
            78,
            121,
        )

        title_style.element.rPr.rFonts.set(
            qn("w:eastAsia"),
            "Arial",
        )

        heading_style = document.styles[
            "Heading 1"
        ]

        heading_style.font.name = "Arial"
        heading_style.font.size = Pt(13)
        heading_style.font.bold = True
        heading_style.font.color.rgb = RGBColor(
            31,
            78,
            121,
        )

        heading_style.element.rPr.rFonts.set(
            qn("w:eastAsia"),
            "Arial",
        )

    @classmethod
    def _add_header(
        cls,
        *,
        document: DocxDocument,
        bundle: DocumentBundle,
    ) -> None:
        """Add a discreet corporate-style document header."""

        section = document.sections[0]
        header = section.header

        paragraph = header.paragraphs[0]

        paragraph.alignment = (
            WD_ALIGN_PARAGRAPH.RIGHT
        )

        run = paragraph.add_run(
            "Spécification du processus"
        )

        run.font.name = "Arial"
        run.font.size = Pt(8)
        run.font.italic = True
        run.font.color.rgb = RGBColor(
            100,
            100,
            100,
        )

        if bundle.metadata.version:
            version_run = paragraph.add_run(
                f" — {bundle.metadata.version}"
            )

            version_run.font.name = "Arial"
            version_run.font.size = Pt(8)
            version_run.font.italic = True
            version_run.font.color.rgb = RGBColor(
                100,
                100,
                100,
            )

    @classmethod
    def _add_title(
        cls,
        *,
        document: DocxDocument,
        title: str,
    ) -> None:
        """Add the process title and a separator."""

        title_paragraph = document.add_paragraph(
            style="Title"
        )

        title_paragraph.alignment = (
            WD_ALIGN_PARAGRAPH.CENTER
        )

        title_run = title_paragraph.add_run(
            title
        )

        title_run.font.name = "Arial"

        title_paragraph.paragraph_format.space_after = (
            Pt(8)
        )

        subtitle = document.add_paragraph()

        subtitle.alignment = (
            WD_ALIGN_PARAGRAPH.CENTER
        )

        subtitle_run = subtitle.add_run(
            "Description narrative du processus"
        )

        subtitle_run.font.name = "Arial"
        subtitle_run.font.size = Pt(12)
        subtitle_run.font.bold = True
        subtitle_run.font.color.rgb = RGBColor(
            80,
            80,
            80,
        )

        subtitle.paragraph_format.space_after = (
            Pt(18)
        )

        cls._add_horizontal_rule(
            subtitle
        )

    @staticmethod
    def _add_metadata_block(
        *,
        document: DocxDocument,
        bundle: DocumentBundle,
    ) -> None:
        table = document.add_table(
            rows=2,
            cols=2,
        )

        table.style = "Table Grid"

        labels = [
            ("Processus", bundle.metadata.title),
            ("Version", bundle.metadata.version),
        ]

        for row, (label, value) in zip(
            table.rows,
            labels,
            strict=False,
        ):
            row.cells[0].text = label
            row.cells[1].text = value or ""

            row.cells[0].paragraphs[0].runs[0].bold = True

        document.add_paragraph()

    @classmethod
    def _add_narrative_section(
        cls,
        *,
        document: DocxDocument,
        paragraphs: list[str],
    ) -> None:
        """Insert every generated narrative paragraph."""

        for text in paragraphs:
            paragraph = document.add_paragraph()

            paragraph.alignment = (
                WD_ALIGN_PARAGRAPH.JUSTIFY
            )

            paragraph.paragraph_format.line_spacing = (
                1.2
            )

            paragraph.paragraph_format.space_after = (
                Pt(12)
            )

            paragraph.paragraph_format.first_line_indent = (
                Cm(0.6)
            )

            run = paragraph.add_run(
                text.strip()
            )

            run.font.name = "Arial"
            run.font.size = Pt(11)

    @staticmethod
    def _add_horizontal_rule(
        paragraph: Paragraph,
    ) -> None:
        """Add a bottom border to a paragraph."""

        paragraph_properties = (
            paragraph._p.get_or_add_pPr()
        )

        paragraph_borders = (
            paragraph_properties.find(
                qn("w:pBdr")
            )
        )

        if paragraph_borders is None:
            paragraph_borders = OxmlElement(
                "w:pBdr"
            )

            paragraph_properties.append(
                paragraph_borders
            )

        bottom_border = OxmlElement(
            "w:bottom"
        )

        bottom_border.set(
            qn("w:val"),
            "single",
        )
        bottom_border.set(
            qn("w:sz"),
            "6",
        )
        bottom_border.set(
            qn("w:space"),
            "6",
        )
        bottom_border.set(
            qn("w:color"),
            "1F4E79",
        )

        paragraph_borders.append(
            bottom_border
        )
